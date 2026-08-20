# python -m ingestion.build_index

# from __future__ import annotations

# import hashlib
# import json
# import os
# import re
# from dataclasses import dataclass, field
# from pathlib import Path
# from typing import Any, Dict, Iterable, List, Optional

# import chromadb
# import docx
# import fitz
# from docx.document import Document as DocxDocument
# from docx.oxml.table import CT_Tbl
# from docx.oxml.text.paragraph import CT_P
# from docx.table import Table
# from docx.text.paragraph import Paragraph

# import config
# import ollama_client

# COLLECTION_NAME = "sop_chunks"
# TARGET_CHARS = int(os.getenv("TARGET_CHUNK_CHARS", "1200"))
# MAX_CHARS = int(os.getenv("MAX_CHUNK_CHARS", "1800"))
# MIN_CHARS = int(os.getenv("MIN_CHUNK_CHARS", "250"))
# OVERLAP_CHARS = int(os.getenv("CHUNK_OVERLAP_CHARS", "160"))
# IMAGE_DIR = Path(os.getenv("SOP_IMAGE_DIR", str(Path(config.VECTORSTORE_DIR) / "sop_images")))


# @dataclass
# class Block:
#     text: str
#     kind: str
#     level: int = 0
#     page: int = 0
#     images: List[str] = field(default_factory=list)


# @dataclass
# class Topic:
#     document: str
#     path: List[str]
#     blocks: List[Block] = field(default_factory=list)
#     images: List[str] = field(default_factory=list)
#     page_start: int = 0
#     page_end: int = 0

#     @property
#     def name(self) -> str:
#         return self.path[-1] if self.path else "General"

#     @property
#     def parent(self) -> str:
#         return self.path[-2] if len(self.path) > 1 else ""

#     @property
#     def path_text(self) -> str:
#         return " > ".join(self.path) if self.path else "General"


# def clean(text: str) -> str:
#     text = text.replace("\xa0", " ")
#     text = re.sub(r"[ \t]+", " ", text)
#     return re.sub(r"\n{3,}", "\n\n", text).strip()


# def slug(text: str) -> str:
#     return re.sub(r"[^A-Za-z0-9._-]+", "_", text).strip("_")[:70] or "item"


# def make_id(*parts: str) -> str:
#     digest = hashlib.sha1("||".join(parts).encode("utf-8")).hexdigest()[:16]
#     return f"{slug(parts[0])[:30]}_{digest}"


# def is_bullet(text: str) -> bool:
#     return bool(re.match(r"^\s*(?:[-•●▪◦*]|\d+[.)]|[A-Za-z][.)]|[ivxlcdmIVXLCDM]+[.)])\s+", text))


# def heading_level(text: str, style: str = "") -> int:
#     style = (style or "").lower()
#     match = re.search(r"heading\s*([1-9])", style)
#     if match:
#         return int(match.group(1))
#     if style in {"title", "subtitle"}:
#         return 1

#     text = clean(text)
#     numbered = re.match(r"^(\d+(?:\.\d+){0,4})[.)]?\s+.+", text)
#     if numbered:
#         return min(numbered.group(1).count(".") + 1, 6)

#     letters = [c for c in text if c.isalpha()]
#     if letters and len(text) <= 100 and len(text.split()) <= 12:
#         uppercase_ratio = sum(c.isupper() for c in letters) / len(letters)
#         if uppercase_ratio >= 0.85:
#             return 2

#     if len(text) <= 110 and text.endswith(":") and len(text.split()) <= 14 and not is_bullet(text):
#         return 3
#     return 0


# def iter_docx_items(document: DocxDocument) -> Iterable[Paragraph | Table]:
#     for child in document.element.body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)


# def save_docx_images(document: DocxDocument, paragraph: Paragraph, out_dir: Path, doc_name: str, counter: int):
#     paths: List[str] = []
#     for blip in paragraph._p.xpath(".//a:blip"):
#         rel_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
#         if not rel_id:
#             continue
#         try:
#             part = document.part.related_parts[rel_id]
#             ext = Path(str(part.partname)).suffix or ".png"
#             target = out_dir / f"{slug(doc_name)}_{counter:04d}{ext}"
#             target.write_bytes(part.blob)
#             paths.append(str(target))
#             counter += 1
#         except Exception as exc:
#             print(f"  ⚠️ DOCX image extraction failed: {exc}")
#     return paths, counter


# def extract_docx(path: str, doc_name: str) -> List[Block]:
#     document = docx.Document(path)
#     out_dir = IMAGE_DIR / slug(doc_name)
#     out_dir.mkdir(parents=True, exist_ok=True)
#     blocks: List[Block] = []
#     counter = 1

#     for item in iter_docx_items(document):
#         if isinstance(item, Paragraph):
#             text = clean(item.text)
#             images, counter = save_docx_images(document, item, out_dir, doc_name, counter)
#             style = item.style.name if item.style else ""
#             level = heading_level(text, style)
#             kind = "heading" if level else "bullet" if is_bullet(text) or "list" in style.lower() else "paragraph"
#             if text or images:
#                 blocks.append(Block(text=text, kind=kind, level=level, images=images))
#         else:
#             rows = []
#             for row in item.rows:
#                 values = [clean(cell.text) for cell in row.cells]
#                 if any(values):
#                     rows.append(" | ".join(values))
#             if rows:
#                 blocks.append(Block(text="\n".join(rows), kind="table"))
#     return blocks


# def extract_pdf(path: str, doc_name: str) -> List[Block]:
#     pdf = fitz.open(path)
#     out_dir = IMAGE_DIR / slug(doc_name)
#     out_dir.mkdir(parents=True, exist_ok=True)
#     blocks: List[Block] = []
#     seen_xrefs = set()
#     counter = 1

#     for page_index, page in enumerate(pdf):
#         page_no = page_index + 1
#         page_images: List[str] = []
#         for image in page.get_images(full=True):
#             xref = image[0]
#             if xref in seen_xrefs:
#                 continue
#             seen_xrefs.add(xref)
#             try:
#                 data = pdf.extract_image(xref)
#                 ext = data.get("ext", "png")
#                 target = out_dir / f"{slug(doc_name)}_page_{page_no:04d}_{counter:04d}.{ext}"
#                 target.write_bytes(data["image"])
#                 page_images.append(str(target))
#                 counter += 1
#             except Exception as exc:
#                 print(f"  ⚠️ PDF image extraction failed on page {page_no}: {exc}")

#         lines = [clean(line) for line in (page.get_text("text") or "").splitlines()]
#         lines = [line for line in lines if line]
#         if not lines and page_images:
#             blocks.append(Block(text="", kind="paragraph", page=page_no, images=page_images))
#             continue

#         for index, line in enumerate(lines):
#             level = heading_level(line)
#             kind = "heading" if level else "bullet" if is_bullet(line) else "paragraph"
#             blocks.append(Block(text=line, kind=kind, level=level, page=page_no, images=page_images if index == 0 else []))
#     pdf.close()
#     return blocks


# def extract_txt(path: str) -> List[Block]:
#     blocks: List[Block] = []
#     for raw in Path(path).read_text(encoding="utf-8").splitlines():
#         line = clean(raw)
#         if not line:
#             continue
#         level = heading_level(line)
#         blocks.append(Block(text=line, kind="heading" if level else "bullet" if is_bullet(line) else "paragraph", level=level))
#     return blocks


# def extract_blocks(path: str, doc_name: str) -> List[Block]:
#     ext = Path(path).suffix.lower()
#     if ext == ".docx":
#         return extract_docx(path, doc_name)
#     if ext == ".pdf":
#         return extract_pdf(path, doc_name)
#     if ext == ".txt":
#         return extract_txt(path)
#     return []


# def to_topics(blocks: List[Block], document: str) -> List[Topic]:
#     topics: List[Topic] = []
#     stack: List[str] = []
#     current = Topic(document=document, path=["General"])

#     def flush():
#         if any(block.text for block in current.blocks) or current.images:
#             topics.append(current)

#     for block in blocks:
#         if block.kind == "heading" and block.text:
#             flush()
#             level = max(1, block.level)
#             stack[:] = stack[: level - 1]
#             stack.append(block.text)
#             current = Topic(document=document, path=list(stack), blocks=[block], images=list(block.images), page_start=block.page, page_end=block.page)
#             continue

#         current.blocks.append(block)
#         if block.page:
#             current.page_start = current.page_start or block.page
#             current.page_end = block.page
#         for image in block.images:
#             if image not in current.images:
#                 current.images.append(image)
#     flush()
#     return topics


# def atomic_units(blocks: List[Block]) -> List[str]:
#     units: List[str] = []
#     bullet_group: List[str] = []

#     def flush_bullets():
#         nonlocal bullet_group
#         if bullet_group:
#             units.append("\n".join(bullet_group))
#             bullet_group = []

#     for block in blocks:
#         text = clean(block.text)
#         if not text:
#             continue
#         if block.kind == "bullet":
#             bullet_group.append(text)
#         else:
#             flush_bullets()
#             units.append(f"Table:\n{text}" if block.kind == "table" else text)
#     flush_bullets()
#     return units


# def split_oversized(text: str) -> List[str]:
#     if len(text) <= MAX_CHARS:
#         return [text]
#     if "\n" in text:
#         result, current, size = [], [], 0
#         for line in text.splitlines():
#             if current and size + len(line) + 1 > MAX_CHARS:
#                 result.append("\n".join(current))
#                 current, size = [], 0
#             current.append(line)
#             size += len(line) + 1
#         if current:
#             result.append("\n".join(current))
#         return result

#     sentences = re.split(r"(?<=[.!?])\s+", text)
#     result, current = [], ""
#     for sentence in sentences:
#         candidate = sentence if not current else f"{current} {sentence}"
#         if current and len(candidate) > MAX_CHARS:
#             result.append(current)
#             current = sentence
#         else:
#             current = candidate
#     if current:
#         result.append(current)
#     return result


# def make_child_chunks(topic: Topic) -> List[Dict[str, Any]]:
#     units: List[str] = []
#     for unit in atomic_units(topic.blocks):
#         units.extend(split_oversized(unit))

#     texts: List[str] = []
#     current: List[str] = []
#     size = 0
#     for unit in units:
#         extra = len(unit) + (2 if current else 0)
#         if current and size >= MIN_CHARS and size + extra > TARGET_CHARS:
#             rendered = "\n\n".join(current)
#             texts.append(rendered)
#             overlap = rendered[-OVERLAP_CHARS:].lstrip()
#             current = [f"Previous context: {overlap}", unit] if overlap else [unit]
#             size = sum(map(len, current)) + 2
#         else:
#             current.append(unit)
#             size += extra
#     if current:
#         texts.append("\n\n".join(current))

#     parent_id = make_id(topic.document, topic.path_text, "parent")
#     chunks: List[Dict[str, Any]] = []
#     for index, body in enumerate(texts):
#         searchable = f"Document: {topic.document}\nTopic: {topic.path_text}\n\n{body}".strip()
#         chunks.append({
#             "id": make_id(topic.document, topic.path_text, str(index), body[:180]),
#             "parent_id": parent_id,
#             "text": searchable,
#             "raw_text": body,
#             "document_name": topic.document,
#             "topic": topic.name,
#             "parent_topic": topic.parent,
#             "topic_path": topic.path_text,
#             "chunk_index": index,
#             "chunk_count": len(texts),
#             "page_start": topic.page_start,
#             "page_end": topic.page_end,
#             "image_paths": topic.images,
#             "has_images": bool(topic.images),
#         })

#     for index, chunk in enumerate(chunks):
#         chunk["previous_chunk_id"] = chunks[index - 1]["id"] if index else ""
#         chunk["next_chunk_id"] = chunks[index + 1]["id"] if index + 1 < len(chunks) else ""
#     return chunks


# def build_parent(topic: Topic) -> Dict[str, Any]:
#     return {
#         "id": make_id(topic.document, topic.path_text, "parent"),
#         "document_name": topic.document,
#         "topic": topic.name,
#         "parent_topic": topic.parent,
#         "topic_path": topic.path_text,
#         "text": "\n\n".join(atomic_units(topic.blocks)),
#         "page_start": topic.page_start,
#         "page_end": topic.page_end,
#         "image_paths": topic.images,
#         "has_images": bool(topic.images),
#     }


# def metadata(chunk: Dict[str, Any]) -> Dict[str, Any]:
#     return {
#         "document_name": chunk["document_name"],
#         "topic": chunk["topic"],
#         "parent_topic": chunk["parent_topic"],
#         "topic_path": chunk["topic_path"],
#         "header": chunk["header"],
#         "sub_header": chunk["sub_header"],
#         "page": int(chunk["page"]),
#         "parent_id": chunk["parent_id"],
#         "chunk_index": int(chunk["chunk_index"]),
#         "chunk_count": int(chunk["chunk_count"]),
#         "previous_chunk_id": chunk["previous_chunk_id"],
#         "next_chunk_id": chunk["next_chunk_id"],
#         "page_start": int(chunk["page_start"] or 0),
#         "page_end": int(chunk["page_end"] or 0),
#         "has_images": bool(chunk["has_images"]),
#         "image_paths_json": json.dumps(chunk["image_paths"], ensure_ascii=False),
#     }


# def build_index() -> None:
#     docs_dir = Path(config.DOCS_DIR)
#     vector_dir = Path(config.VECTORSTORE_DIR)
#     vector_dir.mkdir(parents=True, exist_ok=True)
#     IMAGE_DIR.mkdir(parents=True, exist_ok=True)

#     if not docs_dir.exists():
#         print(f"❌ Documents directory not found: {docs_dir}")
#         return

#     client = chromadb.PersistentClient(path=str(vector_dir))
#     try:
#         client.delete_collection(COLLECTION_NAME)
#         print("🗑️ Deleted existing collection")
#     except Exception:
#         pass
#     collection = client.create_collection(COLLECTION_NAME, metadata={"hnsw:space": "cosine"})

#     all_chunks: List[Dict[str, Any]] = []
#     parents: List[Dict[str, Any]] = []

#     for file_path in sorted(docs_dir.iterdir()):
#         if not file_path.is_file() or file_path.suffix.lower() not in {".docx", ".pdf", ".txt"}:
#             continue
#         print(f"📄 Processing: {file_path.name}")
#         try:
#             blocks = extract_blocks(str(file_path), file_path.stem)
#             topics = to_topics(blocks, file_path.stem)
#             document_chunks: List[Dict[str, Any]] = []
#             for topic in topics:
#                 parents.append(build_parent(topic))
#                 document_chunks.extend(make_child_chunks(topic))
#             all_chunks.extend(document_chunks)
#             image_count = len({img for topic in topics for img in topic.images})
#             print(f"  ✅ Topics: {len(topics)} | Search chunks: {len(document_chunks)} | Images: {image_count}")
#         except Exception as exc:
#             print(f"  ❌ Error processing {file_path.name}: {exc}")
#             import traceback
#             traceback.print_exc()

#     if not all_chunks:
#         print("❌ No chunks created")
#         return

#     print(f"\n🔍 Generating embeddings for {len(all_chunks)} topic-aware chunks...")
#     texts = [chunk["text"] for chunk in all_chunks]
#     embeddings: List[List[float]] = []
#     batch_size = int(os.getenv("EMBED_BATCH_SIZE", "10"))
#     total = (len(texts) + batch_size - 1) // batch_size
#     for start in range(0, len(texts), batch_size):
#         print(f"  Processing batch {start // batch_size + 1}/{total}")
#         try:
#             result = ollama_client.embed_batch(texts[start:start + batch_size])
#             if len(result) != len(texts[start:start + batch_size]):
#                 raise ValueError("Embedding count mismatch")
#             embeddings.extend(result)
#         except Exception as exc:
#             print(f"  ❌ Error generating embeddings: {exc}")
#             return

#     print("📊 Adding to ChromaDB...")
#     collection.add(
#         ids=[chunk["id"] for chunk in all_chunks],
#         documents=texts,
#         embeddings=embeddings,
#         metadatas=[metadata(chunk) for chunk in all_chunks],
#     )

#     (vector_dir / "docstore.json").write_text(json.dumps(all_chunks, indent=2, ensure_ascii=False), encoding="utf-8")
#     (vector_dir / "parentstore.json").write_text(json.dumps(parents, indent=2, ensure_ascii=False), encoding="utf-8")

#     print("\n✅ Index built successfully")
#     print(f"   Parent topics: {len(parents)}")
#     print(f"   Search chunks: {len(all_chunks)}")
#     print(f"   Extracted images: {len({img for chunk in all_chunks for img in chunk['image_paths']})}")
#     print(f"   Image folder: {IMAGE_DIR}")


# if __name__ == "__main__":
#     build_index()

# from __future__ import annotations

# import hashlib
# import json
# import os
# import re
# from dataclasses import dataclass, field
# from pathlib import Path
# from typing import Any, Dict, Iterable, List, Optional

# import chromadb
# import docx
# import fitz
# from docx.document import Document as DocxDocument
# from docx.oxml.table import CT_Tbl
# from docx.oxml.text.paragraph import CT_P
# from docx.table import Table
# from docx.text.paragraph import Paragraph

# import config
# import ollama_client

# COLLECTION_NAME = "sop_chunks"
# TARGET_CHARS = int(os.getenv("TARGET_CHUNK_CHARS", "1200"))
# MAX_CHARS = int(os.getenv("MAX_CHUNK_CHARS", "1800"))
# MIN_CHARS = int(os.getenv("MIN_CHUNK_CHARS", "250"))
# OVERLAP_CHARS = int(os.getenv("CHUNK_OVERLAP_CHARS", "160"))
# IMAGE_DIR = Path(os.getenv("SOP_IMAGE_DIR", str(Path(config.VECTORSTORE_DIR) / "sop_images")))


# @dataclass
# class Block:
#     text: str
#     kind: str
#     level: int = 0
#     page: int = 0
#     images: List[str] = field(default_factory=list)


# @dataclass
# class Topic:
#     document: str
#     path: List[str]
#     blocks: List[Block] = field(default_factory=list)
#     images: List[str] = field(default_factory=list)
#     page_start: int = 0
#     page_end: int = 0

#     @property
#     def name(self) -> str:
#         return self.path[-1] if self.path else "General"

#     @property
#     def parent(self) -> str:
#         return self.path[-2] if len(self.path) > 1 else ""

#     @property
#     def path_text(self) -> str:
#         return " > ".join(self.path) if self.path else "General"


# def clean(text: str) -> str:
#     text = text.replace("\xa0", " ")
#     text = re.sub(r"[ \t]+", " ", text)
#     return re.sub(r"\n{3,}", "\n\n", text).strip()


# def slug(text: str) -> str:
#     return re.sub(r"[^A-Za-z0-9._-]+", "_", text).strip("_")[:70] or "item"


# def make_id(*parts: str) -> str:
#     digest = hashlib.sha1("||".join(parts).encode("utf-8")).hexdigest()[:16]
#     return f"{slug(parts[0])[:30]}_{digest}"


# def is_bullet(text: str) -> bool:
#     return bool(re.match(r"^\s*(?:[-•●▪◦*]|\d+[.)]|[A-Za-z][.)]|[ivxlcdmIVXLCDM]+[.)])\s+", text))


# def heading_level(text: str, style: str = "") -> int:
#     style = (style or "").lower()
#     match = re.search(r"heading\s*([1-9])", style)
#     if match:
#         return int(match.group(1))
#     if style in {"title", "subtitle"}:
#         return 1

#     text = clean(text)
#     numbered = re.match(r"^(\d+(?:\.\d+){0,4})[.)]?\s+.+", text)
#     if numbered:
#         return min(numbered.group(1).count(".") + 1, 6)

#     letters = [c for c in text if c.isalpha()]
#     if letters and len(text) <= 100 and len(text.split()) <= 12:
#         uppercase_ratio = sum(c.isupper() for c in letters) / len(letters)
#         if uppercase_ratio >= 0.85:
#             return 2

#     if len(text) <= 110 and text.endswith(":") and len(text.split()) <= 14 and not is_bullet(text):
#         return 3
#     return 0


# def iter_docx_items(document: DocxDocument) -> Iterable[Paragraph | Table]:
#     for child in document.element.body.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, document)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, document)


# def save_docx_images(document: DocxDocument, paragraph: Paragraph, out_dir: Path, doc_name: str, counter: int):
#     paths: List[str] = []
#     for blip in paragraph._p.xpath(".//a:blip"):
#         rel_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
#         if not rel_id:
#             continue
#         try:
#             part = document.part.related_parts[rel_id]
#             ext = Path(str(part.partname)).suffix or ".png"
#             target = out_dir / f"{slug(doc_name)}_{counter:04d}{ext}"
#             target.write_bytes(part.blob)
#             paths.append(str(target))
#             counter += 1
#         except Exception as exc:
#             print(f"  ⚠️ DOCX image extraction failed: {exc}")
#     return paths, counter


# def extract_docx(path: str, doc_name: str) -> List[Block]:
#     document = docx.Document(path)
#     out_dir = IMAGE_DIR / slug(doc_name)
#     out_dir.mkdir(parents=True, exist_ok=True)
#     blocks: List[Block] = []
#     counter = 1

#     for item in iter_docx_items(document):
#         if isinstance(item, Paragraph):
#             text = clean(item.text)
#             images, counter = save_docx_images(document, item, out_dir, doc_name, counter)
#             style = item.style.name if item.style else ""
#             level = heading_level(text, style)
#             kind = "heading" if level else "bullet" if is_bullet(text) or "list" in style.lower() else "paragraph"
#             if text or images:
#                 blocks.append(Block(text=text, kind=kind, level=level, images=images))
#         else:
#             rows = []
#             for row in item.rows:
#                 values = [clean(cell.text) for cell in row.cells]
#                 if any(values):
#                     rows.append(" | ".join(values))
#             if rows:
#                 blocks.append(Block(text="\n".join(rows), kind="table"))
#     return blocks


# def extract_pdf(path: str, doc_name: str) -> List[Block]:
#     pdf = fitz.open(path)
#     out_dir = IMAGE_DIR / slug(doc_name)
#     out_dir.mkdir(parents=True, exist_ok=True)
#     blocks: List[Block] = []
#     seen_xrefs = set()
#     counter = 1

#     for page_index, page in enumerate(pdf):
#         page_no = page_index + 1
#         page_images: List[str] = []
#         for image in page.get_images(full=True):
#             xref = image[0]
#             if xref in seen_xrefs:
#                 continue
#             seen_xrefs.add(xref)
#             try:
#                 data = pdf.extract_image(xref)
#                 ext = data.get("ext", "png")
#                 target = out_dir / f"{slug(doc_name)}_page_{page_no:04d}_{counter:04d}.{ext}"
#                 target.write_bytes(data["image"])
#                 page_images.append(str(target))
#                 counter += 1
#             except Exception as exc:
#                 print(f"  ⚠️ PDF image extraction failed on page {page_no}: {exc}")

#         lines = [clean(line) for line in (page.get_text("text") or "").splitlines()]
#         lines = [line for line in lines if line]
#         if not lines and page_images:
#             blocks.append(Block(text="", kind="paragraph", page=page_no, images=page_images))
#             continue

#         for index, line in enumerate(lines):
#             level = heading_level(line)
#             kind = "heading" if level else "bullet" if is_bullet(line) else "paragraph"
#             blocks.append(Block(text=line, kind=kind, level=level, page=page_no, images=page_images if index == 0 else []))
#     pdf.close()
#     return blocks


# def extract_txt(path: str) -> List[Block]:
#     blocks: List[Block] = []
#     for raw in Path(path).read_text(encoding="utf-8").splitlines():
#         line = clean(raw)
#         if not line:
#             continue
#         level = heading_level(line)
#         blocks.append(Block(text=line, kind="heading" if level else "bullet" if is_bullet(line) else "paragraph", level=level))
#     return blocks


# def extract_blocks(path: str, doc_name: str) -> List[Block]:
#     ext = Path(path).suffix.lower()
#     if ext == ".docx":
#         return extract_docx(path, doc_name)
#     if ext == ".pdf":
#         return extract_pdf(path, doc_name)
#     if ext == ".txt":
#         return extract_txt(path)
#     return []


# def to_topics(blocks: List[Block], document: str) -> List[Topic]:
#     topics: List[Topic] = []
#     stack: List[str] = []
#     current = Topic(document=document, path=["General"])

#     def flush():
#         if any(block.text for block in current.blocks) or current.images:
#             topics.append(current)

#     for block in blocks:
#         if block.kind == "heading" and block.text:
#             flush()
#             level = max(1, block.level)
#             stack[:] = stack[: level - 1]
#             stack.append(block.text)
#             current = Topic(document=document, path=list(stack), blocks=[block], images=list(block.images), page_start=block.page, page_end=block.page)
#             continue

#         current.blocks.append(block)
#         if block.page:
#             current.page_start = current.page_start or block.page
#             current.page_end = block.page
#         for image in block.images:
#             if image not in current.images:
#                 current.images.append(image)
#     flush()
#     return topics


# def atomic_units(blocks: List[Block]) -> List[str]:
#     units: List[str] = []
#     bullet_group: List[str] = []

#     def flush_bullets():
#         nonlocal bullet_group
#         if bullet_group:
#             units.append("\n".join(bullet_group))
#             bullet_group = []

#     for block in blocks:
#         text = clean(block.text)
#         if not text:
#             continue
#         if block.kind == "bullet":
#             bullet_group.append(text)
#         else:
#             flush_bullets()
#             units.append(f"Table:\n{text}" if block.kind == "table" else text)
#     flush_bullets()
#     return units


# def split_oversized(text: str) -> List[str]:
#     if len(text) <= MAX_CHARS:
#         return [text]
#     if "\n" in text:
#         result, current, size = [], [], 0
#         for line in text.splitlines():
#             if current and size + len(line) + 1 > MAX_CHARS:
#                 result.append("\n".join(current))
#                 current, size = [], 0
#             current.append(line)
#             size += len(line) + 1
#         if current:
#             result.append("\n".join(current))
#         return result

#     sentences = re.split(r"(?<=[.!?])\s+", text)
#     result, current = [], ""
#     for sentence in sentences:
#         candidate = sentence if not current else f"{current} {sentence}"
#         if current and len(candidate) > MAX_CHARS:
#             result.append(current)
#             current = sentence
#         else:
#             current = candidate
#     if current:
#         result.append(current)
#     return result


# def make_child_chunks(topic: Topic, topic_index: int) -> List[Dict[str, Any]]:
#     units: List[str] = []
#     for unit in atomic_units(topic.blocks):
#         units.extend(split_oversized(unit))

#     texts: List[str] = []
#     current: List[str] = []
#     size = 0
#     for unit in units:
#         extra = len(unit) + (2 if current else 0)
#         if current and size >= MIN_CHARS and size + extra > TARGET_CHARS:
#             rendered = "\n\n".join(current)
#             texts.append(rendered)
#             overlap = rendered[-OVERLAP_CHARS:].lstrip()
#             current = [f"Previous context: {overlap}", unit] if overlap else [unit]
#             size = sum(map(len, current)) + 2
#         else:
#             current.append(unit)
#             size += extra
#     if current:
#         texts.append("\n\n".join(current))

#     parent_id = make_id(topic.document, topic.path_text, str(topic_index), "parent")
#     chunks: List[Dict[str, Any]] = []
#     for index, body in enumerate(texts):
#         searchable = f"Document: {topic.document}\nTopic: {topic.path_text}\n\n{body}".strip()
#         chunks.append({
#             "id": make_id(topic.document, topic.path_text, str(topic_index), str(index), body[:180]),
#             "parent_id": parent_id,
#             "text": searchable,
#             "raw_text": body,
#             "document_name": topic.document,
#             "topic": topic.name,
#             "parent_topic": topic.parent,
#             "topic_path": topic.path_text,

#             # Backward-compatible fields used by the existing chatbot.
#             "header": topic.name or "General",
#             "sub_header": topic.parent or "",
#             "page": str(topic.page_start or ""),

#             "chunk_index": index,
#             "chunk_count": len(texts),
#             "page_start": topic.page_start,
#             "page_end": topic.page_end,
#             "image_paths": topic.images,
#             "has_images": bool(topic.images),
#         })

#     for index, chunk in enumerate(chunks):
#         chunk["previous_chunk_id"] = chunks[index - 1]["id"] if index else ""
#         chunk["next_chunk_id"] = chunks[index + 1]["id"] if index + 1 < len(chunks) else ""
#     return chunks


# def build_parent(topic: Topic, topic_index: int) -> Dict[str, Any]:
#     return {
#         "id": make_id(topic.document, topic.path_text, str(topic_index), "parent"),
#         "document_name": topic.document,
#         "topic": topic.name,
#         "parent_topic": topic.parent,
#         "topic_path": topic.path_text,

#         # Backward-compatible fields used by source display and FAQ logic.
#         "header": topic.name or "General",
#         "sub_header": topic.parent or "",
#         "page": str(topic.page_start or ""),

#         "text": "\n\n".join(atomic_units(topic.blocks)),
#         "page_start": topic.page_start,
#         "page_end": topic.page_end,
#         "image_paths": topic.images,
#         "has_images": bool(topic.images),
#     }


# def metadata(chunk: Dict[str, Any]) -> Dict[str, Any]:
#     """
#     Return Chroma-compatible scalar metadata.

#     .get() fallbacks keep the indexer compatible with both old and new
#     chunk dictionaries and prevent KeyError during indexing.
#     """
#     return {
#         "document_name": chunk.get("document_name", ""),
#         "topic": chunk.get("topic", chunk.get("header", "General")),
#         "parent_topic": chunk.get("parent_topic", chunk.get("sub_header", "")),
#         "topic_path": chunk.get(
#             "topic_path",
#             chunk.get("header", chunk.get("topic", "General")),
#         ),

#         # Existing chatbot compatibility.
#         "header": chunk.get("header", chunk.get("topic", "General")),
#         "sub_header": chunk.get(
#             "sub_header",
#             chunk.get("parent_topic", ""),
#         ),
#         "page": str(
#             chunk.get(
#                 "page",
#                 chunk.get("page_start", ""),
#             )
#             or ""
#         ),
#         "parent_id": chunk.get("parent_id", ""),
#         "chunk_index": int(chunk.get("chunk_index", 0)),
#         "chunk_count": int(chunk.get("chunk_count", 1)),
#         "previous_chunk_id": chunk.get("previous_chunk_id", ""),
#         "next_chunk_id": chunk.get("next_chunk_id", ""),
#         "page_start": int(chunk.get("page_start") or 0),
#         "page_end": int(chunk.get("page_end") or 0),
#         "has_images": bool(chunk.get("has_images", False)),
#         "image_paths_json": json.dumps(chunk.get("image_paths", []), ensure_ascii=False),
#     }


# def build_index() -> None:
#     docs_dir = Path(config.DOCS_DIR)
#     vector_dir = Path(config.VECTORSTORE_DIR)
#     vector_dir.mkdir(parents=True, exist_ok=True)
#     IMAGE_DIR.mkdir(parents=True, exist_ok=True)

#     if not docs_dir.exists():
#         print(f"❌ Documents directory not found: {docs_dir}")
#         return

#     client = chromadb.PersistentClient(path=str(vector_dir))
#     try:
#         client.delete_collection(COLLECTION_NAME)
#         print("🗑️ Deleted existing collection")
#     except Exception:
#         pass
#     collection = client.create_collection(COLLECTION_NAME, metadata={"hnsw:space": "cosine"})

#     all_chunks: List[Dict[str, Any]] = []
#     parents: List[Dict[str, Any]] = []

#     for file_path in sorted(docs_dir.iterdir()):
#         if not file_path.is_file() or file_path.suffix.lower() not in {".docx", ".pdf", ".txt"}:
#             continue
#         print(f"📄 Processing: {file_path.name}")
#         try:
#             blocks = extract_blocks(str(file_path), file_path.stem)
#             topics = to_topics(blocks, file_path.stem)
#             document_chunks: List[Dict[str, Any]] = []
#             for topic_index, topic in enumerate(topics):
#                 parents.append(build_parent(topic, topic_index))
#                 document_chunks.extend(make_child_chunks(topic, topic_index))
#             all_chunks.extend(document_chunks)
#             image_count = len({img for topic in topics for img in topic.images})
#             print(f"  ✅ Topics: {len(topics)} | Search chunks: {len(document_chunks)} | Images: {image_count}")
#         except Exception as exc:
#             print(f"  ❌ Error processing {file_path.name}: {exc}")
#             import traceback
#             traceback.print_exc()

#     if not all_chunks:
#         print("❌ No chunks created")
#         return

#     # ChromaDB requires globally unique IDs. Repeated headings/topics are common
#     # in SOPs, so topic_index is included in every parent and child ID.
#     chunk_ids = [chunk["id"] for chunk in all_chunks]
#     parent_ids = [parent["id"] for parent in parents]
#     if len(chunk_ids) != len(set(chunk_ids)):
#         raise ValueError("Duplicate child chunk IDs remain after topic indexing")
#     if len(parent_ids) != len(set(parent_ids)):
#         raise ValueError("Duplicate parent topic IDs remain after topic indexing")

#     print(f"\n🔍 Generating embeddings for {len(all_chunks)} topic-aware chunks...")
#     texts = [chunk["text"] for chunk in all_chunks]
#     embeddings: List[List[float]] = []
#     batch_size = int(os.getenv("EMBED_BATCH_SIZE", "10"))
#     total = (len(texts) + batch_size - 1) // batch_size
#     for start in range(0, len(texts), batch_size):
#         print(f"  Processing batch {start // batch_size + 1}/{total}")
#         try:
#             result = ollama_client.embed_batch(texts[start:start + batch_size])
#             if len(result) != len(texts[start:start + batch_size]):
#                 raise ValueError("Embedding count mismatch")
#             embeddings.extend(result)
#         except Exception as exc:
#             print(f"  ❌ Error generating embeddings: {exc}")
#             return

#     print("📊 Adding to ChromaDB...")
#     collection.add(
#         ids=[chunk["id"] for chunk in all_chunks],
#         documents=texts,
#         embeddings=embeddings,
#         metadatas=[metadata(chunk) for chunk in all_chunks],
#     )

#     (vector_dir / "docstore.json").write_text(json.dumps(all_chunks, indent=2, ensure_ascii=False), encoding="utf-8")
#     (vector_dir / "parentstore.json").write_text(json.dumps(parents, indent=2, ensure_ascii=False), encoding="utf-8")

#     print("\n✅ Index built successfully")
#     print(f"   Parent topics: {len(parents)}")
#     print(f"   Search chunks: {len(all_chunks)}")
#     print(f"   Extracted images: {len({img for chunk in all_chunks for img in chunk['image_paths']})}")
#     print(f"   Image folder: {IMAGE_DIR}")


# if __name__ == "__main__":
#     build_index()


from __future__ import annotations

import hashlib
import json
import os
import shutil
import subprocess
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

import chromadb
import docx
import fitz
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

import config
import ollama_client

COLLECTION_NAME = "sop_chunks"
TARGET_CHARS = int(os.getenv("TARGET_CHUNK_CHARS", "1200"))
MAX_CHARS = int(os.getenv("MAX_CHUNK_CHARS", "1800"))
MIN_CHARS = int(os.getenv("MIN_CHUNK_CHARS", "250"))
# OVERLAP_CHARS = int(os.getenv("CHUNK_OVERLAP_CHARS", "160"))
OVERLAP_CHARS = int(os.getenv("CHUNK_OVERLAP_CHARS", "0"))
IMAGE_DIR = Path(os.getenv("SOP_IMAGE_DIR", str(Path(config.VECTORSTORE_DIR) / "sop_images")))

ACTION_VERBS = {
    "add", "attach", "call", "check", "choose", "click", "complete", "confirm",
    "contact", "continue", "copy", "create", "download", "edit", "email", "enter",
    "forward", "go", "locate", "mail", "modify", "open", "pend", "print",
    "process", "review", "save", "scroll", "select", "send", "status", "submit",
    "trigger", "type", "update", "upload", "verify", "wait"
}

PAGE_NUMBER_RE = re.compile(r"^\s*\d{1,3}\s*$")
MD_HEADING_RE = re.compile(r"^\s*(#{1,6})\s+(.*\S)\s*$")
EMPHASIS_WRAP_RE = re.compile(r"^\s*(\*{1,3}|_{1,3})(.+?)\1\s*$")
SEPARATOR_RE = re.compile(r"^\s*(?:[-_=*]\s*){3,}$")
SYMBOL_BULLET_RE = re.compile(r"^\s*[-\u2022\u25cf\u25aa\u25e6*+]\s+")
INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
SECTION_NUMBER_RE = re.compile(r"^(\d+(?:\.\d+){0,4})[.)]?\s+")

FIELD_LABELS = {
    "purpose", "procedure", "procedures", "overview", "scope", "summary", "note",
    "notes", "example", "examples", "key information", "when to use", "when created",
    "required status", "pre call", "during call", "post call", "result", "reminder",
    "important", "tip", "tips", "definition", "definitions", "exception", "exceptions",
    "steps", "process", "background", "verbiage", "script", "email verbiage",
    "sample", "details", "question",
}

DOC_TITLE_KEYWORDS = (
    "region", "procedures", "taskbook", "training manual", "manual", "handbook",
    "sop", "standard operating", "frequently asked questions", "faq",
)

IMAGE_RASTER_DPI = int(os.getenv("IMAGE_RASTER_DPI", "200"))
PARENT_TOPIC_LEVEL = int(os.getenv("PARENT_TOPIC_LEVEL", str(getattr(config, "PARENT_TOPIC_LEVEL", 3))))
STRIP_MARKDOWN_INLINE = bool(getattr(config, "STRIP_MARKDOWN_INLINE", True))


def normalize_heading_text(text: str):
    """Strip literal markdown markers. Returns (text, md_level, emphasis_width)."""
    text = clean(text)
    md_level = 0
    match = MD_HEADING_RE.match(text)
    if match:
        md_level = len(match.group(1))
        text = match.group(2).strip()

    emphasis_width = 0
    for _ in range(2):
        wrapped = EMPHASIS_WRAP_RE.match(text)
        if not wrapped:
            break
        emphasis_width = max(emphasis_width, len(wrapped.group(1)))
        text = wrapped.group(2).strip()

    return clean(text), md_level, emphasis_width


def section_number_of(text: str) -> str:
    """Return '2.10' from '2.10 Offsite Processing', else ''."""
    match = SECTION_NUMBER_RE.match(normalize_heading_text(text)[0])
    return match.group(1) if match else ""


def is_word_like(core: str) -> bool:
    """
    True for lines made of words, False for codes / amounts / numbers.
    Rejects 'O/S 1234567-89', 'D: 630-285-4037', 'AT&T', 'IAO $125.00', 'CD'.
    """
    letters = [c for c in core if c.isalpha()]
    compact = core.replace(" ", "")
    return bool(compact) and len(letters) >= 4 and (len(letters) / len(compact)) >= 0.5



@dataclass
class Block:
    text: str
    kind: str
    level: int = 0
    page: int = 0
    images: List[str] = field(default_factory=list)


@dataclass
class Topic:
    document: str
    path: List[str]
    blocks: List[Block] = field(default_factory=list)
    images: List[str] = field(default_factory=list)
    page_start: int = 0
    page_end: int = 0

    @property
    def name(self) -> str:
        return self.path[-1] if self.path else "General"

    @property
    def parent(self) -> str:
        return self.path[-2] if len(self.path) > 1 else ""

    @property
    def path_text(self) -> str:
        return " > ".join(self.path) if self.path else "General"


def clean(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def slug(text: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", text).strip("_")[:70] or "item"


def make_id(*parts: str) -> str:
    digest = hashlib.sha1("||".join(parts).encode("utf-8")).hexdigest()[:16]
    return f"{slug(parts[0])[:30]}_{digest}"


def is_bullet(text: str) -> bool:
    return bool(re.match(r"^\s*(?:[-•●▪◦*]|\d+[.)]|[A-Za-z][.)]|[ivxlcdmIVXLCDM]+[.)])\s+", text))


def looks_like_instruction(text: str) -> bool:
    """Return True when a short line looks like an action step, not a heading."""
    cleaned = clean(text).strip(" \"'")
    words = re.findall(r"[A-Za-z]+", cleaned)
    if not words:
        return False
    return words[0].lower() in ACTION_VERBS


# def heading_level(text: str, style: str = "", *, trust_style: bool = False) -> int:
#     """Detect genuine headings conservatively."""
#     style = (style or "").lower()

#     if trust_style:
#         match = re.search(r"heading\s*([1-9])", style)
#         if match:
#             return int(match.group(1))
#         if style in {"title", "subtitle"}:
#             return 1

#     text = clean(text)
#     if not text or PAGE_NUMBER_RE.match(text):
#         return 0

#     numbered = re.match(r"^(\d+(?:\.\d+){0,4})[.)]?\s+(.+)$", text)
#     if numbered:
#         number, title = numbered.groups()
#         title = title.strip()
#         if looks_like_instruction(title):
#             return 0
#         if len(title) > 120 or len(title.split()) > 16:
#             return 0
#         if title.endswith((".", "?", "!", ";", ",")):
#             return 0
#         return min(number.count(".") + 1, 6)

#     letters = [c for c in text if c.isalpha()]
#     if letters and len(text) <= 100 and len(text.split()) <= 12:
#         uppercase_ratio = sum(c.isupper() for c in letters) / len(letters)
#         if uppercase_ratio >= 0.90 and not looks_like_instruction(text):
#             return 2

#     return 0

def heading_level(
    text: str,
    style: str = "",
    *,
    trust_style: bool = False,
    is_bold: bool = False,
) -> int:
    """
    Detect genuine SOP headings.

    Levels: 1 = document title, 2 = numbered section, 3 = numbered subsection,
            4 = bold in-section label (never starts a new topic).

    Accepted:
        HARTFORD REGION PROCEDURES              -> 1
        ### **WEST REGION **                    -> 1
        20. ROI'S / **20. ROI'S**               -> 2
        1. Central - Facility Will Not Provide  -> 2
        2.10 Offsite Processing                 -> 3
        Step 9 - Apply the Final Status         -> 4
        Required Documents / Image Requirement  -> 4

    Rejected (these were splitting procedures in half):
        O/S 1234567-89, D: 630-285-4037, AT&T, CD, IAO $125.00
        - DMRS, - MOD, - VRC, • IEHP
        26. Open a new email to confirm your signature
        Purpose / Procedure / Question
    """
    style_lower = (style or "").strip().lower()
    core, md_level, emphasis_width = normalize_heading_text(text)
    md_emphasis = emphasis_width in (1, 2)

    if not core or PAGE_NUMBER_RE.fullmatch(core) or SEPARATOR_RE.match(core):
        return 0

    # A bulleted line is a list item, never a heading.
    if SYMBOL_BULLET_RE.match(core):
        return 0

    if trust_style:
        style_match = re.search(r"heading\s*([1-9])", style_lower)
        if style_match:
            return int(style_match.group(1))
        if style_lower in {"title", "subtitle"}:
            return 1
        if "list" in style_lower:
            return 0

    if md_level:
        if md_level <= 2 or any(k in core.lower() for k in DOC_TITLE_KEYWORDS):
            return 1
        return 2

    emphasised = bool(is_bold or md_emphasis)

    numbered_match = re.match(r"^(\d+(?:\.\d+){0,4})[.)]?\s+(.+)$", core)
    if numbered_match:
        number, title = numbered_match.groups()
        title = title.strip(" *_")
        number_depth = number.count(".") + 1

        if not title or len(title) > 140 or len(title.split()) > 20:
            return 0

        # "2.10", "12.4.1" - always a real subsection.
        if "." in number:
            return min(number_depth + 1, 6)

        # ALL-CAPS titles are headings, not steps. Several SOP section titles
        # begin with a word that is also an action verb ("23. STATUS LETTERS",
        # "27. PROCESSING OFFSITES"), and the instruction check was rejecting
        # them - so their content was swallowed by the previous section.
        title_letters = [c for c in title if c.isalpha()]
        title_is_caps = (
            bool(title_letters)
            and sum(c.isupper() for c in title_letters) / len(title_letters) >= 0.9
        )

        if not title_is_caps and looks_like_instruction(title):
            return 0
        if title.endswith((".", "!", ";", ",")):
            return 0

        if emphasised:
            return 2
        if any(k in title.lower() for k in (
            "region", "procedures", "frequently asked questions",
            "highlights and exceptions", "training manual", "table of contents",
        )):
            return 2
        if not trust_style and len(title.split()) <= 12:
            return 2
        return 0

    # All-caps headings.
    letters = [c for c in core if c.isalpha()]
    if is_word_like(core) and len(core) <= 100 and len(core.split()) <= 12:
        uppercase_ratio = sum(c.isupper() for c in letters) / len(letters)
        if uppercase_ratio >= 0.90 and not looks_like_instruction(core):
            # A trailing colon marks a lead-in label introducing the list below
            # it ("EXAMPLES WHEN A STATUS LETTER IS NEEDED:"), not a new
            # section. Level 4 keeps that list inside its parent section.
            if core.rstrip().endswith(":"):
                return 4
            return 1 if any(k in core.lower() for k in DOC_TITLE_KEYWORDS) else 2

    # Bold in-section labels. Level 4 so they never split a procedure.
    if (
        emphasised
        and emphasis_width != 3
        and is_word_like(core)
        and 3 <= len(core) <= 80
        and len(core.split()) <= 10
        and core.strip(" :").lower() not in FIELD_LABELS
        and (core[0].isupper() or core[0].isdigit())
        and not is_bullet(core)
        and not looks_like_instruction(core)
        and not core.endswith((".", "!", ";", ",", "?"))
        and "=" not in core
    ):
        return 4

    return 0

def normalize_pdf_line(line: str, page_no: int) -> str:
    """Remove page numbers and repair body text joined to a page number."""
    line = clean(line)
    if not line or PAGE_NUMBER_RE.match(line):
        return ""

    match = re.match(rf"^{page_no}[.)]?\s+(.+)$", line)
    if match:
        remainder = match.group(1).strip()
        if looks_like_instruction(remainder):
            return remainder

    return line


def iter_docx_items(document: DocxDocument) -> Iterable[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)

def is_docx_numbered_list(paragraph: Paragraph) -> bool:
    """
    Detect Word numbered or bulleted list paragraphs.

    This prevents numbered statuses and numbered procedure steps from
    being incorrectly treated as SOP section headings.
    """
    try:
        paragraph_properties = paragraph._p.pPr
        return (
            paragraph_properties is not None
            and paragraph_properties.numPr is not None
        )
    except Exception:
        return False


WEB_SAFE_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
CONVERTIBLE_IMAGE_EXTS = {".emf", ".wmf", ".tif", ".tiff"}


def to_web_safe_image(source: Path) -> Optional[Path]:
    """
    Convert a metafile (EMF/WMF/TIFF) into a PNG a browser can display.

    Word stores pasted screenshots as EMF: 89 of the 99 images in West.docx are
    EMF, which no browser can render - the user sees a broken image and the alt
    text instead of the picture. Pillow can rasterise EMF/WMF on Windows.

    Returns the PNG path, or None if conversion is not possible (in which case
    the image is skipped rather than served as an unrenderable file).
    """
    suffix = source.suffix.lower()
    if suffix in WEB_SAFE_IMAGE_EXTS:
        return source
    if suffix not in CONVERTIBLE_IMAGE_EXTS:
        return None

    # No "reuse existing target" shortcut: image numbering shifts whenever a
    # document changes, so a leftover PNG from an earlier build would be
    # silently reused for a completely different picture.
    target = source.with_suffix(".png")
    target.unlink(missing_ok=True)

    if (
        _convert_with_pillow(source, target)
        or _convert_with_windows_gdi(source, target)
        or _convert_with_soffice(source, target)
    ):
        source.unlink(missing_ok=True)
        return target

    target.unlink(missing_ok=True)
    return None


def _convert_with_pillow(source: Path, target: Path) -> bool:
    """Pillow rasterises WMF/EMF through the Windows GDI. Windows only."""
    try:
        from PIL import Image
    except ImportError:
        return False
    try:
        with Image.open(source) as image:
            try:
                image.load(dpi=IMAGE_RASTER_DPI)   # WmfImagePlugin dpi hint
            except TypeError:
                image.load()
            image.convert("RGB").save(target, "PNG", optimize=True)
        return target.exists() and target.stat().st_size > 0
    except Exception:
        return False


def _convert_with_soffice(source: Path, target: Path) -> bool:
    """
    Fallback for Linux/Docker: convert with LibreOffice headless.

    Pillow can only rasterise EMF/WMF through the Windows GDI, so inside a
    Linux container every EMF would be skipped. Install libreoffice-draw in
    the image (or set INSTALL_IMAGE_TOOLS=true when building) to enable this.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False
    try:
        result = subprocess.run(
            [soffice, "--headless", "--norestore", "--convert-to", "png",
             "--outdir", str(target.parent), str(source)],
            capture_output=True, timeout=120,
        )
        produced = target.parent / (source.stem + ".png")
        if result.returncode == 0 and produced.exists() and produced.stat().st_size > 0:
            if produced != target:
                produced.replace(target)
            return True
    except Exception:
        pass
    return False


def _convert_with_windows_gdi(source: Path, target: Path) -> bool:
    """
    Fallback for Windows: rasterise the metafile with .NET System.Drawing.
    Used when Pillow is missing or its WMF loader is unavailable.
    """
    if os.name != "nt":
        return False

    scale = max(1.0, IMAGE_RASTER_DPI / 96.0)
    script = (
        "Add-Type -AssemblyName System.Drawing; "
        f"$mf = New-Object System.Drawing.Imaging.Metafile '{source}'; "
        "$h = $mf.GetMetafileHeader(); "
        f"$w = [int]($mf.Width * {scale}); $ht = [int]($mf.Height * {scale}); "
        "if ($w -lt 1) {{ $w = $mf.Width }}; if ($ht -lt 1) {{ $ht = $mf.Height }}; "
        "$bmp = New-Object System.Drawing.Bitmap $w, $ht; "
        "$g = [System.Drawing.Graphics]::FromImage($bmp); "
        "$g.Clear([System.Drawing.Color]::White); "
        "$g.DrawImage($mf, 0, 0, $w, $ht); "
        f"$bmp.Save('{target}', [System.Drawing.Imaging.ImageFormat]::Png); "
        "$g.Dispose(); $bmp.Dispose(); $mf.Dispose();"
    ).replace("{{", "{").replace("}}", "}")

    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
            capture_output=True, timeout=60,
        )
        if result.returncode != 0:
            return False
        return target.exists() and target.stat().st_size > 0
    except Exception:
        return False


def save_docx_images(document: DocxDocument, paragraph: Paragraph, out_dir: Path, doc_name: str, counter: int):
    paths: List[str] = []
    for blip in paragraph._p.xpath(".//a:blip"):
        rel_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not rel_id:
            continue
        try:
            part = document.part.related_parts[rel_id]
            ext = Path(str(part.partname)).suffix or ".png"
            target = out_dir / f"{slug(doc_name)}_{counter:04d}{ext}"
            target.write_bytes(part.blob)

            usable = to_web_safe_image(target)
            counter += 1
            if usable is None:
                target.unlink(missing_ok=True)
                continue

            # Store a path RELATIVE to the image root so the index stays valid
            # if the project is moved or deployed to a different directory.
            paths.append(str(usable.relative_to(IMAGE_DIR)).replace(os.sep, "/"))
        except Exception as exc:
            print(f"  ⚠️ DOCX image extraction failed: {exc}")
    return paths, counter


# def extract_docx(path: str, doc_name: str) -> List[Block]:
#     document = docx.Document(path)
#     out_dir = IMAGE_DIR / slug(doc_name)
#     out_dir.mkdir(parents=True, exist_ok=True)
#     blocks: List[Block] = []
#     counter = 1

#     for item in iter_docx_items(document):
#         if isinstance(item, Paragraph):
#             text = clean(item.text)
#             images, counter = save_docx_images(document, item, out_dir, doc_name, counter)
#             style = item.style.name if item.style else ""
#             level = heading_level(text, style, trust_style=True)
#             kind = "heading" if level else "bullet" if is_bullet(text) or "list" in style.lower() else "paragraph"
#             if text or images:
#                 blocks.append(Block(text=text, kind=kind, level=level, images=images))
#         else:
#             rows = []
#             for row in item.rows:
#                 values = [clean(cell.text) for cell in row.cells]
#                 if any(values):
#                     rows.append(" | ".join(values))
#             if rows:
#                 blocks.append(Block(text="\n".join(rows), kind="table"))
#     return blocks

def extract_docx(path: str, doc_name: str) -> List[Block]:
    document = docx.Document(path)
    out_dir = IMAGE_DIR / slug(doc_name)
    # Wipe this document's images first. Files from a previous build otherwise
    # survive, and because image numbering shifts when the document changes a
    # stale "West_0055.png" gets reused for a different image entirely - which
    # is how a non-compliance status diagram ended up under
    # "36. MAINTENANCE FORM".
    if out_dir.exists():
        for stale in out_dir.iterdir():
            if stale.is_file():
                stale.unlink(missing_ok=True)
    out_dir.mkdir(parents=True, exist_ok=True)
    blocks: List[Block] = []
    counter = 1

    for item in iter_docx_items(document):
        if isinstance(item, Paragraph):
            raw_text = clean(item.text)
            images, counter = save_docx_images(document, item, out_dir, doc_name, counter)
            style = item.style.name if item.style else ""
            style_lower = style.lower()

            # Horizontal rules ("______", "---") are layout, not content.
            if raw_text and SEPARATOR_RE.match(raw_text):
                if images:
                    blocks.append(Block(text="", kind="paragraph", images=images))
                continue

            # Every paragraph in these SOPs uses the 'Normal' Word style, so
            # bold runs are the only heading signal available.
            runs = [run for run in item.runs if run.text.strip()]
            is_bold = bool(runs) and all(bool(run.bold) for run in runs)

            is_list_paragraph = (
                "list" in style_lower
                or is_docx_numbered_list(item)
            )

            if is_list_paragraph:
                level = 0
            else:
                level = heading_level(
                    raw_text,
                    style,
                    trust_style=True,
                    is_bold=is_bold,
                )

            if level:
                text = normalize_heading_text(raw_text)[0]
                kind = "heading"
            else:
                text = INLINE_BOLD_RE.sub(r"\1", raw_text) if STRIP_MARKDOWN_INLINE else raw_text
                kind = "bullet" if (is_list_paragraph or is_bullet(text)) else "paragraph"

            if text or images:
                blocks.append(Block(text=text, kind=kind, level=level, images=images))
        else:
            rows = []
            for row in item.rows:
                values = [clean(cell.text) for cell in row.cells]
                if any(values):
                    rows.append(" | ".join(values))
            if rows:
                blocks.append(Block(text="\n".join(rows), kind="table"))
    return blocks


def extract_pdf(path: str, doc_name: str) -> List[Block]:
    pdf = fitz.open(path)
    out_dir = IMAGE_DIR / slug(doc_name)
    out_dir.mkdir(parents=True, exist_ok=True)
    blocks: List[Block] = []
    seen_xrefs = set()
    counter = 1

    for page_index, page in enumerate(pdf):
        page_no = page_index + 1
        page_images: List[str] = []
        for image in page.get_images(full=True):
            xref = image[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            try:
                data = pdf.extract_image(xref)
                ext = data.get("ext", "png")
                target = out_dir / f"{slug(doc_name)}_page_{page_no:04d}_{counter:04d}.{ext}"
                target.write_bytes(data["image"])
                page_images.append(str(target))
                counter += 1
            except Exception as exc:
                print(f"  ⚠️ PDF image extraction failed on page {page_no}: {exc}")

        lines = [normalize_pdf_line(line, page_no) for line in (page.get_text("text") or "").splitlines()]
        lines = [line for line in lines if line]
        if not lines and page_images:
            blocks.append(Block(text="", kind="paragraph", page=page_no, images=page_images))
            continue

        for index, line in enumerate(lines):
            level = heading_level(line)
            kind = "heading" if level else "bullet" if is_bullet(line) else "paragraph"
            blocks.append(Block(text=line, kind=kind, level=level, page=page_no, images=page_images if index == 0 else []))
    pdf.close()
    return blocks


def extract_txt(path: str) -> List[Block]:
    blocks: List[Block] = []
    for raw in Path(path).read_text(encoding="utf-8").splitlines():
        line = clean(raw)
        if not line:
            continue
        level = heading_level(line)
        blocks.append(Block(text=line, kind="heading" if level else "bullet" if is_bullet(line) else "paragraph", level=level))
    return blocks


def extract_blocks(path: str, doc_name: str) -> List[Block]:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return extract_docx(path, doc_name)
    if ext == ".pdf":
        return extract_pdf(path, doc_name)
    if ext == ".txt":
        return extract_txt(path)
    return []


def to_topics(blocks: List[Block], document: str) -> List[Topic]:
    topics: List[Topic] = []
    stack: List[str] = []
    current = Topic(document=document, path=["General"])

    def flush():
        if any(block.text for block in current.blocks) or current.images:
            topics.append(current)

    for block in blocks:
        if block.kind == "heading" and block.text:
            level = max(1, block.level)

            # Headings deeper than PARENT_TOPIC_LEVEL are in-section labels
            # ("Step 9 - ...", "Required Documents", "Image Requirement").
            # They stay inside the topic so a procedure is never cut in half.
            if level > PARENT_TOPIC_LEVEL:
                current.blocks.append(
                    Block(text=block.text, kind="subheading", level=level, images=block.images)
                )
                for image in block.images:
                    if image not in current.images:
                        current.images.append(image)
                continue

            flush()
            stack[:] = stack[: level - 1]
            stack.append(block.text)
            current = Topic(document=document, path=list(stack), blocks=[], images=list(block.images), page_start=block.page, page_end=block.page)
            continue

        current.blocks.append(block)
        if block.page:
            current.page_start = current.page_start or block.page
            current.page_end = block.page
        for image in block.images:
            if image not in current.images:
                current.images.append(image)
    flush()
    return topics


def atomic_units_with_images(blocks: List[Block]):
    """
    Same grouping as atomic_units(), but each unit carries the images that
    appeared alongside it, so a chunk only gets the pictures belonging to the
    text it actually contains rather than every image in the whole topic.
    """
    units = []
    bullet_group: List[str] = []
    bullet_images: List[str] = []
    pending_images: List[str] = []

    def flush_bullets():
        nonlocal bullet_group, bullet_images
        if bullet_group:
            units.append(("\n".join(bullet_group), list(bullet_images)))
            bullet_group = []
            bullet_images = []

    for block in blocks:
        text = clean(block.text)
        images = list(block.images or [])

        if not text:
            # Image-only paragraph: hold it for the next unit of text.
            pending_images.extend(images)
            continue

        images = pending_images + images
        pending_images = []

        if block.kind == "bullet":
            bullet_group.append(text)
            bullet_images.extend(images)
        else:
            flush_bullets()
            rendered = f"Table:\n{text}" if block.kind == "table" else text
            units.append((rendered, images))
    flush_bullets()

    if pending_images:
        if units:
            units[-1][1].extend(pending_images)
        else:
            units.append(("", pending_images))

    return [(text, imgs) for text, imgs in units if text or imgs]


def atomic_units(blocks: List[Block]) -> List[str]:
    """Backward-compatible text-only view (used for parent topic bodies)."""
    return [text for text, _ in atomic_units_with_images(blocks) if text]


def split_oversized(text: str) -> List[str]:
    if len(text) <= MAX_CHARS:
        return [text]
    if "\n" in text:
        result, current, size = [], [], 0
        for line in text.splitlines():
            if current and size + len(line) + 1 > MAX_CHARS:
                result.append("\n".join(current))
                current, size = [], 0
            current.append(line)
            size += len(line) + 1
        if current:
            result.append("\n".join(current))
        return result

    sentences = re.split(r"(?<=[.!?])\s+", text)
    result, current = [], ""
    for sentence in sentences:
        candidate = sentence if not current else f"{current} {sentence}"
        if current and len(candidate) > MAX_CHARS:
            result.append(current)
            current = sentence
        else:
            current = candidate
    if current:
        result.append(current)
    return result


def make_child_chunks(topic: Topic, topic_index: int) -> List[Dict[str, Any]]:
    # (text, images) pairs, so each chunk keeps only its own pictures.
    units = []
    for unit_text, unit_images in atomic_units_with_images(topic.blocks):
        fragments = split_oversized(unit_text) if unit_text else [""]
        for position, fragment in enumerate(fragments):
            units.append((fragment, list(unit_images) if position == 0 else []))

    texts: List[str] = []
    chunk_images: List[List[str]] = []
    current: List[str] = []
    current_images: List[str] = []
    size = 0
    for unit, images in units:
        if not unit:
            current_images.extend(images)
            continue
        extra = len(unit) + (2 if current else 0)
        if current and size >= MIN_CHARS and size + extra > TARGET_CHARS:
            rendered = "\n\n".join(current)
            texts.append(rendered)
            chunk_images.append(list(dict.fromkeys(current_images)))

            overlap_units: List[str] = []
            if OVERLAP_CHARS > 0 and current:
                last_unit = current[-1]
                if len(last_unit) <= OVERLAP_CHARS:
                    overlap_units = [last_unit]

            current = overlap_units + [unit]
            current_images = list(images)
            size = sum(map(len, current)) + (2 * max(0, len(current) - 1))
        else:
            current.append(unit)
            current_images.extend(images)
            size += extra
    if current:
        texts.append("\n\n".join(current))
        chunk_images.append(list(dict.fromkeys(current_images)))

    parent_id = make_id(topic.document, topic.path_text, str(topic_index), "parent")
    chunks: List[Dict[str, Any]] = []
    for index, body in enumerate(texts):
        searchable = f"Document: {topic.document}\nTopic: {topic.path_text}\n\n{body}".strip()
        chunks.append({
            "id": make_id(topic.document, topic.path_text, str(topic_index), str(index), body[:180]),
            "parent_id": parent_id,
            "text": searchable,
            "raw_text": body,
            "document_name": topic.document,
            "topic": topic.name,
            "parent_topic": topic.parent,
            "topic_path": topic.path_text,
            "display_title": topic.name or topic.parent or "General",

            # Backward-compatible fields used by the existing chatbot.
            "header": topic.name or "General",
            "sub_header": topic.parent or "",
            "page": str(topic.page_start or ""),

            "chunk_index": index,
            "chunk_count": len(texts),
            "page_start": topic.page_start,
            "page_end": topic.page_end,
            "image_paths": chunk_images[index] if index < len(chunk_images) else [],
            "has_images": bool(chunk_images[index] if index < len(chunk_images) else []),
        })

    for index, chunk in enumerate(chunks):
        chunk["previous_chunk_id"] = chunks[index - 1]["id"] if index else ""
        chunk["next_chunk_id"] = chunks[index + 1]["id"] if index + 1 < len(chunks) else ""
    return chunks


def build_parent(topic: Topic, topic_index: int) -> Dict[str, Any]:
    return {
        "id": make_id(topic.document, topic.path_text, str(topic_index), "parent"),
        "document_name": topic.document,
        "topic": topic.name,
        "parent_topic": topic.parent,
        "topic_path": topic.path_text,
        "display_title": topic.name or topic.parent or "General",

        # Backward-compatible fields used by source display and FAQ logic.
        "header": topic.name or "General",
        "sub_header": topic.parent or "",
        "page": str(topic.page_start or ""),

        "text": "\n\n".join(atomic_units(topic.blocks)),
        "page_start": topic.page_start,
        "page_end": topic.page_end,
        "image_paths": topic.images,
        "has_images": bool(topic.images),
    }


def metadata(chunk: Dict[str, Any]) -> Dict[str, Any]:
    """
    Return Chroma-compatible scalar metadata.

    .get() fallbacks keep the indexer compatible with both old and new
    chunk dictionaries and prevent KeyError during indexing.
    """
    return {
        "document_name": chunk.get("document_name", ""),
        "topic": chunk.get("topic", chunk.get("header", "General")),
        "parent_topic": chunk.get("parent_topic", chunk.get("sub_header", "")),
        "topic_path": chunk.get(
            "topic_path",
            chunk.get("header", chunk.get("topic", "General")),
        ),
        "display_title": chunk.get(
            "display_title",
            chunk.get("topic", chunk.get("header", "General")),
        ),

        # Existing chatbot compatibility.
        "header": chunk.get("header", chunk.get("topic", "General")),
        "sub_header": chunk.get(
            "sub_header",
            chunk.get("parent_topic", ""),
        ),
        "page": str(
            chunk.get(
                "page",
                chunk.get("page_start", ""),
            )
            or ""
        ),
        "parent_id": chunk.get("parent_id", ""),
        "chunk_index": int(chunk.get("chunk_index", 0)),
        "chunk_count": int(chunk.get("chunk_count", 1)),
        "previous_chunk_id": chunk.get("previous_chunk_id", ""),
        "next_chunk_id": chunk.get("next_chunk_id", ""),
        "page_start": int(chunk.get("page_start") or 0),
        "page_end": int(chunk.get("page_end") or 0),
        "has_images": bool(chunk.get("has_images", False)),
        "image_paths_json": json.dumps(chunk.get("image_paths", []), ensure_ascii=False),
    }


def build_index() -> None:
    docs_dir = Path(config.DOCS_DIR)
    vector_dir = Path(config.VECTORSTORE_DIR)
    vector_dir.mkdir(parents=True, exist_ok=True)
    IMAGE_DIR.mkdir(parents=True, exist_ok=True)

    if not docs_dir.exists():
        print(f"❌ Documents directory not found: {docs_dir}")
        return

    client = chromadb.PersistentClient(path=str(vector_dir))
    try:
        client.delete_collection(COLLECTION_NAME)
        print("🗑️ Deleted existing collection")
    except Exception:
        pass
    collection = client.create_collection(COLLECTION_NAME, metadata={"hnsw:space": "cosine"})

    all_chunks: List[Dict[str, Any]] = []
    parents: List[Dict[str, Any]] = []

    for file_path in sorted(docs_dir.iterdir()):
        if not file_path.is_file() or file_path.suffix.lower() not in {".docx", ".pdf", ".txt"}:
            continue
        print(f"📄 Processing: {file_path.name}")
        try:
            blocks = extract_blocks(str(file_path), file_path.stem)
            topics = to_topics(blocks, file_path.stem)
            document_chunks: List[Dict[str, Any]] = []
            for topic_index, topic in enumerate(topics):
                parents.append(build_parent(topic, topic_index))
                document_chunks.extend(make_child_chunks(topic, topic_index))
            all_chunks.extend(document_chunks)
            image_count = len({img for topic in topics for img in topic.images})
            print(f"  ✅ Topics: {len(topics)} | Search chunks: {len(document_chunks)} | Images: {image_count}")
        except Exception as exc:
            print(f"  ❌ Error processing {file_path.name}: {exc}")
            import traceback
            traceback.print_exc()

    if not all_chunks:
        print("❌ No chunks created")
        return

    # ChromaDB requires globally unique IDs. Repeated headings/topics are common
    # in SOPs, so topic_index is included in every parent and child ID.
    chunk_ids = [chunk["id"] for chunk in all_chunks]
    parent_ids = [parent["id"] for parent in parents]
    if len(chunk_ids) != len(set(chunk_ids)):
        raise ValueError("Duplicate child chunk IDs remain after topic indexing")
    if len(parent_ids) != len(set(parent_ids)):
        raise ValueError("Duplicate parent topic IDs remain after topic indexing")

    print(f"\n🔍 Generating embeddings for {len(all_chunks)} topic-aware chunks...")
    texts = [chunk["text"] for chunk in all_chunks]
    embeddings: List[List[float]] = []
    batch_size = int(os.getenv("EMBED_BATCH_SIZE", "10"))
    total = (len(texts) + batch_size - 1) // batch_size
    for start in range(0, len(texts), batch_size):
        print(f"  Processing batch {start // batch_size + 1}/{total}")
        try:
            result = ollama_client.embed_batch(texts[start:start + batch_size])
            if len(result) != len(texts[start:start + batch_size]):
                raise ValueError("Embedding count mismatch")
            embeddings.extend(result)
        except Exception as exc:
            print(f"  ❌ Error generating embeddings: {exc}")
            return

    print("📊 Adding to ChromaDB...")
    collection.add(
        ids=[chunk["id"] for chunk in all_chunks],
        documents=texts,
        embeddings=embeddings,
        metadatas=[metadata(chunk) for chunk in all_chunks],
    )

    (vector_dir / "docstore.json").write_text(json.dumps(all_chunks, indent=2, ensure_ascii=False), encoding="utf-8")
    (vector_dir / "parentstore.json").write_text(json.dumps(parents, indent=2, ensure_ascii=False), encoding="utf-8")

    print("\n✅ Index built successfully")
    print(f"   Parent topics: {len(parents)}")
    print(f"   Search chunks: {len(all_chunks)}")
    print(f"   Extracted images: {len({img for chunk in all_chunks for img in chunk['image_paths']})}")
    print(f"   Image folder: {IMAGE_DIR}")


if __name__ == "__main__":
    build_index()