"""
Readers that pull text out of .docx and .pdf files while preserving
header/section structure, which is the key signal we need later to
tell visually-similar SOPs apart.
"""
import os
import re
from docx import Document
import fitz  # PyMuPDF


def read_docx(filepath):
    """
    Returns a list of dicts: {header, sub_header, page, text}
    Headers are detected from Word paragraph styles (Heading 1 / Heading 2).
    If a doc has no styled headings, falls back to bold/all-caps line heuristics.
    """
    doc = Document(filepath)
    blocks = []
    current_h1 = "General"
    current_h2 = ""
    buffer = []

    def flush():
        text = "\n".join(buffer).strip()
        if text:
            blocks.append({
                "header": current_h1,
                "sub_header": current_h2,
                "page": None,  # docx has no reliable page numbers
                "text": text
            })
        buffer.clear()

    for para in doc.paragraphs:
        style = (para.style.name or "").lower()
        text = para.text.strip()
        if not text:
            continue

        is_h1 = "heading 1" in style or "title" in style
        is_h2 = "heading 2" in style
        # fallback heuristic: short, bold, all-caps-ish line acting as a header
        looks_like_header = (
            not is_h1 and not is_h2
            and len(text) < 80
            and (text.isupper() or all(r.bold for r in para.runs if r.text.strip()))
            and len(para.runs) > 0
        )

        if is_h1:
            flush()
            current_h1 = text
            current_h2 = ""
        elif is_h2 or looks_like_header:
            flush()
            current_h2 = text
        else:
            buffer.append(text)

    flush()

    # also pull text from tables (SOPs often hide key rules in tables)
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows_text.append(" | ".join(cells))
        if rows_text:
            blocks.append({
                "header": current_h1,
                "sub_header": "Table",
                "page": None,
                "text": "\n".join(rows_text)
            })

    return blocks


HEADER_LINE_RE = re.compile(r"^[A-Z][A-Za-z0-9 \-/&]{3,80}$")


def read_pdf(filepath):
    """
    Returns a list of dicts: {header, sub_header, page, text}
    PDFs don't carry style metadata reliably, so headers are detected with a
    heuristic: short lines, larger font size relative to the page's body text.
    """
    blocks = []
    current_h1 = "General"
    buffer = []
    pdf = fitz.open(filepath)

    def flush(page_num):
        text = "\n".join(buffer).strip()
        if text:
            blocks.append({
                "header": current_h1,
                "sub_header": "",
                "page": page_num,
                "text": text
            })
        buffer.clear()

    for page_num, page in enumerate(pdf, start=1):
        page_dict = page.get_text("dict")
        sizes = [
            span["size"]
            for block in page_dict.get("blocks", [])
            for line in block.get("lines", [])
            for span in line.get("spans", [])
        ]
        body_size = max(set(sizes), key=sizes.count) if sizes else 10

        for block in page_dict.get("blocks", []):
            for line in block.get("lines", []):
                line_text = "".join(span["text"] for span in line["spans"]).strip()
                if not line_text:
                    continue
                max_span_size = max((s["size"] for s in line["spans"]), default=body_size)
                is_header = (
                    max_span_size > body_size + 1.5
                    and len(line_text) < 90
                ) or bool(HEADER_LINE_RE.match(line_text) and len(line_text.split()) <= 8)

                if is_header:
                    flush(page_num)
                    current_h1 = line_text
                else:
                    buffer.append(line_text)
        flush(page_num)

    pdf.close()
    return blocks


def read_document(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".docx":
        return read_docx(filepath)
    elif ext == ".pdf":
        return read_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
